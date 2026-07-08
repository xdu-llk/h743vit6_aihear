"""Update Word report with current project data."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches
import re

DOC = r'C:\Users\jeveux\Desktop\基于STM32H7的边缘AI语音识别居家安防系统.docx'
doc = Document(DOC)

# Build paragraph index map
paras = {}
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        paras[i] = text

# Find key sections by content matching
def find_section(keyword):
    for idx, text in paras.items():
        if keyword in text:
            return idx
    return -1

# ===== Replace: 1.5 嵌入式软件设计 =====
old_1_5 = (
    '软件基于手动集成的 FreeRTOS v10.6.2 内核，采用单任务架构：AudioInferTask 优先级 3、8KB 栈空间。'
)
new_1_5 = (
    '软件基于手动集成的 FreeRTOS v10.6.2 内核，采用单任务架构：AudioInferTask 优先级 3、8KB 栈空间。'
    'DMA 中断通过 vTaskNotifyGiveFromISR 发送通知，任务唤醒后一次性读取 32 帧 PCM 数据（共 5472 采样点，'
    '约 342ms 窗口），帧顺序从最旧到最新（与 Python 训练严格一致，避免 spectrogram 时间轴反转），'
    '经 CMSIS-DSP 512-pt RFFT + Mel 滤波器组生成 40x32 特征矩阵，送入 TFLite Micro DS-CNN 模型进行 3 分类推理。'
)

# Add PCM scaling bug paragraph after 1.5 body
pcm_bug = (
    '预处理管线严格对齐 Python 训练管线，期间解决了两个关键预处理 BUG：\n'
    '(1) PCM 缩放修正——INMP441 输出 24bit I2S 数据，经(int32_t)raw >> 8 转换后范围约 [-8388608, 8388607]。'
    '初期错误除以 32768（16bit范围）导致信号放大 256 倍，特征空间完全错位。'
    '修正为除以 8388608.0f 后 FFT 输入归一化至 [-1,1]，与 Python librosa.load 一致。\n'
    '(2) Spectrogram 时间轴反转——Audio_ReadSamples 按 offset 读取时，offset=512 为最近帧、offset=5472 为最旧帧。'
    '初期 for(f=0->31) 先喂最新帧导致 spec[:,0]=最新帧而 Python lm[:,0]=最早帧，spectrogram 互为镜像。'
    '修正为 for(f=31->0) 从最旧帧开始。'
)

float32_note = (
    '模型采用 float32 格式，TFLite 模型 78KB，arena 大小 512KB（AXI SRAM），环形缓冲区 32KB 置于 DTCM。'
    '推理耗时 < 10ms。未采用 INT8 量化——onnx2tf 的 flatbuffer_direct 模式使用固定 scale=1/128 进行全整数量化，'
    '无法提供校准数据，量化后精度从 92.6% 骤降至约 25%。Keras 重建 PyTorch 权重时 BN 层 set_weights 存在时序问题'
    '（逐层对比验证确认 d1.bn1 开始分叉）。最终采用 float32 部署方案，以 512KB arena 换取可靠性。'
)

alarm_logic = (
    '报警逻辑包括：峰值阈值门控（仅 peak > 50000 触发推理）、margin 置信度过滤'
    '（top1-top2 >= 1.5 才触发告警，滤除低置信度误报）、5 秒冷却防止重复告警。'
    '蜂鸣器 PWM 驱动扫频警笛，RGB LED 根据状态显示多种颜色模式。'
    '按键采用 button.c 独立消抖模块，PC4 上拉输入，连续 3 次读到低电平才确认按下。'
)

# Find and update 1.5
idx_15_start = find_section('1.5 嵌入式软件设计')
if idx_15_start >= 0:
    # Replace 1.5 body
    for p in doc.paragraphs:
        if '软件基于手动集成的 FreeRTOS v10.6.2' in p.text and 'AudioInferTask' in p.text:
            p.text = new_1_5
            break

    # Add PCM bug paragraph
    for i, p in enumerate(doc.paragraphs):
        if '预处理管线严格对齐 Python 训练管线' in p.text and 'Hann 窗' in p.text:
            p.text = pcm_bug
            break

    for i, p in enumerate(doc.paragraphs):
        if 'TFLite 模型经 INT8 全整数量化' in p.text:
            p.text = float32_note
            break

    for i, p in enumerate(doc.paragraphs):
        if '报警逻辑包括' in p.text and '峰值阈值门控' in p.text:
            p.text = alarm_logic
            break

# ===== Replace: 1.6 模型训练 =====
new_1_6 = (
    '数据集经过三次迭代后采用纯语义三类数据集，总计 874 条 16kHz 单声道样本：\n'
    '  • help（救命）：102 条 = PC 录制的真人"救命"语音 + Microsoft Edge TTS 中文"救命"'
    '（6 种音色 × 5 速率 × 5 音调）\n'
    '  • glass_break（玻璃碎裂）：210 条 = PC 录制的敲玻璃声 + ESC-50 glass_breaking 类'
    '（40 源文件 × 8 窗口峰值提取）\n'
    '  • impact（撞击）：562 条 = PC 录制的拍手/敲门声 + ESC-50 door_wood_knock + clapping 类'
    '（各 40 源文件 × 8 窗口）\n\n'
    '数据增强采用在线方式（训练时实时生成）：音高偏移（±3 半音，30% 概率）、时间拉伸（0.85-1.15x，30% 概率）、'
    '高斯噪声叠加（SNR 15-30dB，40% 概率）、音量增益（0.6-1.4x，30% 概率）、'
    'SpecAugment 频域时域 Masking（50% 概率）。\n\n'
    '模型采用 DS-CNN 架构：Conv2D(1->32) + 4 × DepthwiseSeparableConv2d（32->64->128）+ '
    'GlobalAveragePool + FC(128->3)，约 1.9 万参数，36M MACs。训练策略包括 WeightedRandomSampler（类平衡）、'
    'CrossEntropyLoss(class_weight)、MixUp（α=0.2）、AdamW（lr=0.002）、'
    'CosineAnnealingWarmRestarts 学习率调度。训练 80 轮，验证集准确率 92.6%'
    '（help=100%, glass=94.8%, impact=96.8%）。\n\n'
    '模型导出流程：PyTorch → ONNX（torch.onnx.export, TorchScript 后端）→ float32 TFLite'
    '（onnx2tf flatbuffer_direct 转换）→ C 数组（dscnn_model.cc）。未采用 INT8 量化的原因：'
    'onnx2tf 的 flatbuffer_direct 模式不支持校准数据集、Keras 权重传输存在 BN 层 set_weights 时序问题。'
)

for i, p in enumerate(doc.paragraphs):
    if '数据集使用 ESC-50 公开环境声音数据集' in p.text:
        p.text = new_1_6
        break

# Second paragraph of 1.6 (model architecture)
for i, p in enumerate(doc.paragraphs):
    if '模型采用 DS-CNN 架构' in p.text and '1 个标准 Conv2d' in p.text:
        p.text = ''  # already covered in new_1_6
        break

# ===== Replace: 1.9 性能指标 (tables) =====
# Try to find and replace in tables
for t_idx, table in enumerate(doc.tables):
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if '固件大小' in cells[0]:
            # Update cell
            row.cells[1].text = '374KB text + 0.5KB data + 534KB bss'
        elif '模型大小' in cells[0]:
            row.cells[1].text = '78KB float32 TFLite（1.9万参数）'
        elif '模型推理耗时' in cells[0]:
            row.cells[1].text = '< 10ms（480MHz Cortex-M7）'
        elif 'TFLite 验证准确率' in cells[0]:
            row.cells[1].text = '92.6%（3分类，float32）'

# ===== Add 3.1 new content =====
new_3_1 = (
    '本项目最深刻的体会是：AI 模型效果严重依赖训练数据质量和预处理的一致性。以下按时间线记录五个关键踩坑点：\n'
    '(1) 数据质量清洗：初期 ESC-50 数据集经固定偏移切片后产生 372 个静音文件（6.3%），'
    'glass_break 类最严重（15.8%）。改用峰值检测窗口提取后零静音。\n'
    '(2) 类别语义映射：初期将 crying_baby 映射为 help、footsteps 映射为 normal、'
    'fireworks/hand_saw/keyboard_typing 混入 impact——频谱特征与目标类别完全不同。'
    '修正为纯语义映射：help 仅含中文"救命"、glass 仅含 glass_breaking、impact 仅含 door_wood_knock+clapping。\n'
    '(3) 训练-推理预处理一致性（三个致命 BUG）：\n'
    '   a) PCM 缩放范围不匹配——初期除以 32768（16bit）而非 8388608（24bit>>8），信号放大 256 倍\n'
    '   b) Spectrogram 时间轴反转——帧喂入顺序颠倒导致 MCU 与 Python 特征互为镜像\n'
    '   c) float32 vs INT8 量化陷阱——onnx2tf fixed-scale 量化精度崩塌，Keras BN 权重传输失败\n'
    '(4) 域迁移问题：PC 耳机麦克风录制的训练数据在 INMP441 设备端完全无法泛化，'
    '不同麦克风的频率响应差异在 mel 频谱中仍显著，z-score 归一化无法完全补偿。\n'
    '(5) 边际置信度过滤：引入 top1-top2 margin 阈值（>=1.5），滤除模型不确定的边界样本'
)

for i, p in enumerate(doc.paragraphs):
    if '数据驱动的 AI 开发' in p.text:
        p.text = '3.1 数据驱动的 AI 开发'
        break

for i, p in enumerate(doc.paragraphs):
    if '本项目最深刻的体会是' in p.text and 'AI 模型效果严重依赖' in p.text and '初期使用错误的' in p.text:
        p.text = new_3_1
        break

# ===== Replace: 1.8 成果展示 =====
new_1_8 = (
    '系统实现了完整的功能闭环：\n'
    '  • 安静环境：OLED 显示 ARMED + 实时峰值\n'
    '  • 拍桌子/敲门：检测 IMPACT（margin >= 1.5），红色 LED 急闪 + 蜂鸣器警笛，MQTT 推送\n'
    '  • 敲玻璃：检测 GLASS，同上触发全部报警\n'
    '  • 喊"救命"：检测 HELP，同上触发全部报警\n'
    '  • 低置信度声音：显示 [WEAK] + 类别，不触发告警\n'
    '  • 按键切换撤防：OLED 显示 DISARMED，停止检测（软件消抖）\n'
    '  • Web 仪表盘：实时告警时间线，运行时长，设备在线状态\n'
    '  • 微信推送：WxPusher 免费公众号推送告警信息到手机'
)

for i, p in enumerate(doc.paragraphs):
    if '拍桌子/摔东西' in p.text and 'IMPACT' in p.text:
        # This is a bullet list item; update parent
        pass

# Find the 1.8 section and update bullets
for i, p in enumerate(doc.paragraphs):
    if '安静环境' in p.text and 'ARMED' in p.text and '蓝色 LED' in p.text:
        p.text = '  • 安静环境：OLED 显示 ARMED + 实时峰值'
    elif '拍桌子/摔东西' in p.text and 'IMPACT/GLASS' in p.text:
        p.text = '  • 拍桌子/敲门：检测 IMPACT（margin >= 1.5），红色 LED 急闪 + 蜂鸣器警笛，MQTT 推送'
    elif '按键切换撤防' in p.text:
        p.text = '  • 按键切换撤防：OLED 显示 DISARMED，停止检测（软件消抖）'

# ===== Update 3.3 add button + memory =====
for i, p in enumerate(doc.paragraphs):
    if '蜂鸣器调试中' in p.text and 'PWM 输出极性' in p.text:
        p.text = p.text + (
            '\n按键消抖也是硬件调试中的常见问题。初期 freertos.c 中直接用 HAL_GPIO_ReadPin 读取 + '
            'btn_prev/btn_now 边沿检测，因机械抖动导致判断失效。改为 button.c 独立消抖模块——'
            '连续 3 次读到低电平才确认按下。'
            '\n内存布局方面，DTCM 仅 128KB。环形缓冲区（32KB）和 PCM 录音 buffer（64KB）同时存在时会导致 '
            'DTCM 溢出。最终将大型静态 buffer 通过链接脚本 .ring_buf 段放置到 512KB AXI SRAM。'
        )
        break

# ===== Update 3.4 =====
for i, p in enumerate(doc.paragraphs):
    if '深度学习：CNN 模型设计' in p.text and 'INT8 量化' in p.text:
        p.text = '  • 深度学习：CNN 模型设计、数据集工程（清洗→标注→增强→质量审计）、域迁移问题诊断、ONNX→TFLite 部署管线'
        break
    if '工程调试：DWT 计时' in p.text:
        p.text = (
            '  • 工程调试：DWT 计时、HardFault 诊断（SRAM ECC 冷启动）、逻辑分析仪使用、链路逐级排查\n'
            '  • C++/Python 联合调试：逐层对比 PyTorch vs Keras 中间张量定位 BN 权重传输错误、串口 PCM dump 验证 MCU 端数据质量'
        )
        break

# Save
out = r'C:\Users\jeveux\Desktop\基于STM32H7的边缘AI语音识别居家安防系统_更新版.docx'
doc.save(out)
print(f'Saved: {out}')
